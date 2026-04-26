from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE_DIR = Path.cwd()
REPORT_DATE = "2026-04-26"
OUTPUT = BASE_DIR / f"market_briefing_{REPORT_DATE}.docx"
CHECKED_AT = "2026-04-26 16:40 KST"


SOURCES = [
    ("Investing.com 한국 홈/시장", "https://kr.investing.com/"),
    ("Investing.com KOSPI", "https://www.investing.com/indices/kospi"),
    ("Investing.com S&P 500", "https://kr.investing.com/indices/us-spx-500"),
    ("한국경제 - SK하이닉스 1분기 실적", "https://www.hankyung.com/article/2026042327061"),
    ("한국경제 - 한국 1분기 GDP", "https://www.hankyung.com/article/2026042327221"),
    ("한국경제 - 호르무즈 해협 리스크", "https://www.hankyung.com/article/2026042326981"),
    ("한국경제 - 미국 증시 사상 최고치", "https://www.hankyung.com/amp/2026042302726"),
    ("한국경제 - 월스트리트나우", "https://www.hankyung.com/article/202604164147i"),
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def cm_to_twips(value):
    return str(int(value / 2.54 * 1440))


def set_table_width(table, total_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), cm_to_twips(total_cm))


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), cm_to_twips(width_cm))


def style_table(table, header_fill="1F4E79"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Apple SD Gothic Neo"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)


def add_table(doc, headers, rows, widths=None):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(3)
    run = heading.add_run(" | ".join(headers))
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(31, 78, 121)
    run.font.name = "Apple SD Gothic Neo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    for row in rows:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        parts = [f"{headers[i]}: {row[i]}" for i in range(len(headers))]
        r = p.add_run(" | ".join(parts))
        r.font.size = Pt(8.5)
        r.font.name = "Apple SD Gothic Neo"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Apple SD Gothic Neo"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Apple SD Gothic Neo"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
        rest = text[len(bold_prefix):]
        r2 = p.add_run(rest)
        r2.font.name = "Apple SD Gothic Neo"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    else:
        r = p.add_run(text)
        r.font.name = "Apple SD Gothic Neo"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.name = "Apple SD Gothic Neo"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")


def add_callout(doc, title, text):
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(8)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(31, 78, 121)
    r1.font.name = "Apple SD Gothic Neo"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.3)
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(text)
    r2.font.size = Pt(9.5)
    r2.font.name = "Apple SD Gothic Neo"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Apple SD Gothic Neo"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")
    styles["Normal"].font.size = Pt(9.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Apple SD Gothic Neo"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Market Briefing Bot | 정보 브리핑이며 투자 권유가 아닙니다.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 100, 100)


def build_report():
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(42)
    r = title.add_run("아침 증시 브리핑")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)
    r.font.name = "Apple SD Gothic Neo"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Apple SD Gothic Neo")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("한국증시 / 미국증시 / 핵심 이슈 리포트")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(90, 90, 90)

    meta_rows = [
        ("작성 기준일", "2026-04-26 일요일 KST"),
        ("시장 기준", "한국: 2026-04-24 마감, 미국: 2026-04-25 KST 표기 마감"),
        ("확인 시각", CHECKED_AT),
        ("주요 출처", "Investing.com, 한국경제/한경 글로벌마켓"),
    ]
    add_table(doc, ["항목", "내용"], meta_rows, widths=[3.5, 12.5])
    add_callout(
        doc,
        "오늘의 한 줄 결론",
        "정규장이 쉬는 일요일 기준 최신 마감은 AI/반도체가 위험자산 선호를 지탱하지만, 호르무즈 해협 교착과 고유가 가능성이 주초 변동성의 핵심 변수입니다.",
    )

    doc.add_page_break()

    add_heading(doc, "Executive Summary", 1)
    add_bullets(
        doc,
        [
            "한국 증시는 최신 마감 기준 코스피가 6,475.63으로 보합권, 코스닥은 1,203.84로 2.51% 상승했다. 대형주는 차익실현이 있었지만 성장주와 중소형 테마의 탄력이 살아 있었다.",
            "미국 증시는 S&P 500이 7,165.08로 0.80%, 나스닥종합지수가 24,836.60으로 1.63% 상승했다. 다우는 49,230.71로 0.16% 하락해 기술주 우위 장세가 뚜렷했다.",
            "한국경제는 SK하이닉스의 1분기 영업이익률 72%, 매출 52조5763억원, 영업이익 37조6103억원을 보도했다. HBM/eSSD 수요가 한국 증시와 GDP 서프라이즈의 핵심 동력으로 해석된다.",
            "호르무즈 해협은 휴전에도 통행 정상화가 지연되고 있다. 한국경제 보도 기준 하루 통행량은 전쟁 전 130척 수준에서 1척까지 줄었고, 기뢰 제거에는 반년이 걸릴 수 있다는 분석이 있다.",
            "이번 주 체크포인트는 빅테크 실적, 미국 국채 입찰, 일본은행 이벤트, 원/달러 환율과 유가, 그리고 국내 반도체 밸류체인의 차익실현 압력이다.",
        ],
    )

    add_heading(doc, "주요 시장 지표", 1)
    add_table(
        doc,
        ["구분", "지표", "최신값", "등락", "의미"],
        [
            ("한국", "코스피", "6,475.63", "-0.18 / 0.00%", "전고점 부근에서 숨고르기"),
            ("한국", "코스닥", "1,203.84", "+29.53 / +2.51%", "성장주·테마주 수급 강세"),
            ("미국", "다우존스", "49,230.71", "-79.61 / -0.16%", "블루칩은 상대적 부진"),
            ("미국", "S&P 500", "7,165.08", "+56.68 / +0.80%", "대형 성장주 중심 상승"),
            ("미국", "나스닥종합", "24,836.60", "+398.09 / +1.63%", "AI·반도체 모멘텀 지속"),
            ("변동성", "CBOE VIX", "18.71", "-0.60 / -3.11%", "공포 완화, 단 이벤트 리스크 잔존"),
            ("원자재", "브렌트유 선물", "105.33", "+0.26 / +0.25%", "호르무즈 리스크 프리미엄 반영"),
            ("원자재", "WTI 선물", "94.40", "-1.45 / -1.51%", "단기 차익실현과 공급 우려 혼재"),
            ("환율", "달러/원", "1,475.89", "+4.09 / -0.28%", "고환율 부담은 남아 있음"),
        ],
        widths=[2.0, 3.0, 2.7, 3.2, 5.2],
    )

    add_heading(doc, "한국 증시 핵심 흐름", 1)
    add_body(
        doc,
        "코스피는 6,400대 중반에서 마감하며 사상 고점 부근의 부담을 소화하고 있다. Investing.com의 최신 시장표는 코스피 6,475.63, 코스닥 1,203.84를 제시했다. 코스피는 보합권이지만 코스닥은 2%대 상승으로 마감해 대형 반도체의 차익실현과 성장주 순환매가 동시에 나타난 하루로 볼 수 있다.",
    )
    add_body(
        doc,
        "주도 축은 여전히 반도체다. 한국경제는 SK하이닉스가 1분기 매출 52조5763억원, 영업이익 37조6103억원, 영업이익률 72%를 기록했다고 보도했다. HBM3E 공급 확대와 범용 D램 가격 상승, 서버용 eSSD 호조가 복합적으로 작용했다.",
    )
    add_body(
        doc,
        "다만 좋은 뉴스가 곧바로 추가 상승을 보장하지는 않는다. 같은 보도에서 코스피는 장중 6,557.75까지 오르며 고점을 경신한 뒤 차익실현으로 상승폭을 줄였다. 주초에는 반도체 실적 기대가 남아 있는지, 아니면 이미 가격에 상당 부분 반영됐는지를 확인해야 한다.",
    )
    add_table(
        doc,
        ["한국 이슈", "긍정 요인", "경계 요인", "주초 확인 포인트"],
        [
            ("반도체", "HBM·eSSD 수요, SK하이닉스 실적 서프라이즈", "중국 메모리 진입, 고객사 공급망 다변화", "삼성전자·하이닉스 수급, 장비/소재 동반 여부"),
            ("거시", "1분기 GDP 1.7% 깜짝 성장", "반도체 쏠림과 고환율 부담", "수출주 외 내수주 확산 여부"),
            ("수급", "외국인 기술주 관심 지속 가능성", "고점권 차익실현", "장중 고점 돌파 후 종가 유지 여부"),
        ],
        widths=[2.5, 4.5, 4.5, 4.5],
    )

    add_heading(doc, "미국 증시 핵심 흐름", 1)
    add_body(
        doc,
        "미국은 기술주가 지수를 다시 끌어올렸다. Investing.com 기준 S&P 500은 7,165.08로 0.80% 상승, 나스닥종합지수는 24,836.60으로 1.63% 상승했다. 다우는 0.16% 하락해 시장의 폭이 넓다기보다 AI·반도체·대형 성장주 중심으로 압축된 장세였다.",
    )
    add_body(
        doc,
        "한국경제 글로벌 기사 흐름도 같은 방향이다. 휴전 연장과 기업 호실적 기대 속에 S&P 500과 나스닥이 사상 최고치를 경신했고, 테크가 시장을 더 끌어올릴 수 있다는 월가 기대가 보도됐다. 다만 동일가중 지수와 산업주가 상대적으로 약하면 상승의 질은 계속 점검해야 한다.",
    )
    add_body(
        doc,
        "주초 미국 시장은 빅테크 실적과 금리 이벤트를 동시에 본다. AI 투자 지속성이 확인되면 고PER 부담을 완화할 수 있지만, 자본지출이 커지는 만큼 수익성 가이던스가 약하면 나스닥 중심으로 변동성이 커질 수 있다.",
    )
    add_table(
        doc,
        ["미국 이슈", "시장 반응", "왜 중요한가", "리스크"],
        [
            ("AI/빅테크", "나스닥 1.63% 상승", "지수 상승의 핵심 엔진", "실적 기대 미달 시 쏠림 해소"),
            ("중동/유가", "VIX 하락에도 유가 프리미엄 지속", "인플레이션과 마진에 영향", "호르무즈 교착 장기화"),
            ("금리", "10년물·입찰 일정 주목", "성장주 할인율에 직접 영향", "입찰 부진 또는 매파 발언"),
        ],
        widths=[2.8, 4.0, 4.4, 4.4],
    )

    add_heading(doc, "업종 및 테마별 이슈", 1)
    add_body(
        doc,
        "반도체/HBM: 한국과 미국 모두 AI 인프라 투자에 연결된 반도체 체인이 주도주다. 한국은 SK하이닉스의 이익률 서프라이즈가 밸류체인 기대를 키웠고, 미국은 엔비디아·AMD·인텔 등 반도체 관련 거래가 지수에 영향을 주는 구간이다.",
    )
    add_body(
        doc,
        "소프트웨어/클라우드: 한국경제 월스트리트나우는 하이퍼스케일러 지출과 빅테크 실적 기대를 언급했다. 하드웨어 랠리 이후 소프트웨어가 따라붙는지, 또는 AI 투자 부담으로 마진 우려가 커지는지가 관건이다.",
    )
    add_body(
        doc,
        "에너지/방산/해운: 호르무즈 통행 제한은 유가, 운임, 보험료, 정유·화학 마진에 동시에 영향을 준다. 유가 상승은 에너지주에는 긍정적일 수 있지만, 한국처럼 에너지 수입 의존도가 높은 시장에는 비용 상승 요인이다.",
    )
    add_table(
        doc,
        ["테마", "관찰 대상", "상방 시나리오", "하방 시나리오"],
        [
            ("HBM", "SK하이닉스, 삼성전자, 장비/소재", "고부가 메모리 가격 강세 지속", "중국 진입·고객사 단가 압박"),
            ("AI 인프라", "엔비디아, AMD, 클라우드 기업", "빅테크 CAPEX 확대 가이던스", "수익화 지연과 밸류에이션 부담"),
            ("에너지", "브렌트/WTI, 정유·화학", "호르무즈 장기 교착", "협상 진전과 위험 프리미엄 축소"),
            ("환율", "USD/KRW, 수출주·수입주", "수출 대형주 이익 방어", "내수 비용 부담과 외국인 수급 둔화"),
        ],
        widths=[2.3, 4.2, 4.7, 4.7],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "매크로, 금리, 환율, 원자재 체크", 1)
    add_body(
        doc,
        "한국경제는 한국 1분기 GDP가 1.7% 깜짝 증가했고 5년 6개월 만에 가장 높은 성장률이라고 보도했다. 반도체 호황이 중동 전쟁 충격을 눌렀다는 해석이지만, 성장 동력이 특정 산업에 집중된 점은 경기의 폭을 점검해야 한다는 뜻이기도 하다.",
    )
    add_body(
        doc,
        "환율은 여전히 부담이다. Investing.com과 한국경제의 환율 표기는 원/달러가 1,470원대 중반에 머무는 것으로 나타난다. 고환율은 수출주에는 일부 완충 효과가 있지만, 에너지·원자재 수입 비용과 외국인 자금 흐름에는 부담이다.",
    )
    add_body(
        doc,
        "원자재에서는 호르무즈가 핵심이다. 한국경제는 해협 통행량이 전쟁 전 하루 130척 수준에서 1척까지 줄었고, 기뢰 제거에는 반년이 걸릴 수 있다는 분석을 전했다. 주초 유가가 추가로 튀면 인플레이션 기대와 국채금리, 성장주 밸류에이션이 동시에 흔들릴 수 있다.",
    )

    add_heading(doc, "오늘 주목할 일정과 리스크", 1)
    add_table(
        doc,
        ["구분", "체크포인트", "브리핑 관점"],
        [
            ("경제 이벤트", "미국 2년물·5년물 국채 입찰", "수요 부진 시 금리 상승, 성장주 부담"),
            ("중앙은행", "일본은행 통화정책성명서·금리결정", "엔화·원화·아시아 수급에 파급 가능"),
            ("기업 실적", "미국 빅테크 실적 시즌", "AI CAPEX와 마진 가이던스 확인"),
            ("지정학", "호르무즈 해협 통행 정상화 여부", "유가·운임·인플레이션 기대의 핵심"),
            ("국내 수급", "반도체 대형주 차익실현 강도", "코스피 고점 돌파 지속성 판단"),
        ],
        widths=[3.0, 5.8, 7.0],
    )

    add_heading(doc, "투자자가 볼 체크포인트", 1)
    add_bullets(
        doc,
        [
            "코스피가 장중 고점권을 다시 돌파하더라도 종가로 버티는지 확인한다.",
            "SK하이닉스 호실적이 삼성전자, 장비, 소재, 후공정으로 확산되는지 본다.",
            "나스닥 상승이 소수 빅테크에만 갇히는지, 소프트웨어와 중소형 성장주로 넓어지는지 확인한다.",
            "브렌트유가 100달러대 중반에서 더 오르면 한국 기업 마진과 환율에 부담이 커질 수 있다.",
            "원/달러 1,470원대가 안정되는지, 외국인 순매수와 함께 내려오는지 점검한다.",
        ],
    )
    add_callout(
        doc,
        "브리핑 톤",
        "이 문서는 시장 정보 요약이며 특정 종목 또는 자산의 매수·매도 권유가 아닙니다. 휴장일 데이터는 최신 마감 기준이며, 일부 Investing.com 수치는 지연 시세 또는 CFD 표기가 섞일 수 있습니다.",
    )

    add_heading(doc, "출처 목록", 1)
    for idx, (name, url) in enumerate(SOURCES, start=1):
        add_body(doc, f"{idx}. {name}: {url}")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build_report()
    print(out)
